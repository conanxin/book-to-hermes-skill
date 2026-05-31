#!/usr/bin/env python3
"""
Extract text from a document for book-to-hermes-skill processing.

Supported formats:
  - Markdown / plain text — direct read
  - PDF — pdftotext → PyPDF2 → pdfminer.six fallback chain
  - HTML — bs4 + lxml (local files only, no network, no script execution)
  - DOCX — python-docx (local files only, no macros, no embedded objects)
  - EPUB — ebooklib (local files only, no network, no script execution)

Safety rules:
  - NO automatic package installation
  - NO network requests
  - NO arbitrary shell execution
  - Missing dependencies → clear error, exit 1

Output: JSON to stdout with extraction result.
"""

import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

TEXT_EXTENSIONS = {".txt", ".text", ".md", ".markdown"}
SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".docx", ".epub", *TEXT_EXTENSIONS}

# Optional venv for isolated dependencies
VENV_PYTHON = os.path.expanduser("~/.hermes/venvs/book-to-hermes-skill/bin/python")


def read_text_file(path: str) -> str | None:
    """Read a text file with encoding fallback."""
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "cp1252", "latin-1"):
        try:
            return Path(path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception:
            return None
    return None


def extract_with_pdftotext(pdf_path: str) -> str | None:
    if not shutil.which("pdftotext"):
        return None
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", pdf_path, "-"],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass
    return None


def extract_with_pypdf2(pdf_path: str) -> str | None:
    try:
        import PyPDF2
        text_parts = []
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    text_parts.append("")
        return "\n".join(text_parts)
    except ImportError:
        return None
    except Exception:
        return None


def extract_with_pdfminer(pdf_path: str) -> str | None:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(pdf_path)
    except ImportError:
        return None
    except Exception:
        return None


def extract_pdf(pdf_path: str) -> tuple[str, str, list[str]]:
    """Extract text from PDF with fallback chain."""
    warnings = []

    text = extract_with_pdftotext(pdf_path)
    if text and text.strip():
        return text, "pdftotext", warnings
    warnings.append("pdftotext not available or failed")

    text = extract_with_pypdf2(pdf_path)
    if text and text.strip():
        return text, "PyPDF2", warnings
    warnings.append("PyPDF2 not installed")

    text = extract_with_pdfminer(pdf_path)
    if text and text.strip():
        return text, "pdfminer.six", warnings
    warnings.append("pdfminer.six not installed")

    return "", "", warnings + ["ERROR: Could not extract text from PDF"]


def extract_html(html_path: str) -> tuple[str, str, list[str]]:
    """Extract text from local HTML file using BeautifulSoup."""
    warnings = []
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return "", "", ["ERROR: beautifulsoup4 not installed. Install: pip install beautifulsoup4 lxml"]

    text = read_text_file(html_path)
    if text is None:
        return "", "", ["ERROR: Could not read HTML file"]

    # Use lxml if available, fallback to html.parser
    try:
        soup = BeautifulSoup(text, "lxml")
        parser = "lxml"
    except Exception:
        soup = BeautifulSoup(text, "html.parser")
        parser = "html.parser"
        warnings.append("lxml not available, using html.parser fallback")

    # Remove noise elements (no script execution, no network)
    for tag_name in ["script", "style", "noscript", "svg", "nav", "footer", "aside", "iframe", "embed", "object"]:
        for tag in soup.find_all(tag_name):
            tag.decompose()

    # Try to get title
    title = ""
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.get_text(strip=True)

    # Extract content in document order
    lines = []
    if title:
        lines.append(f"# {title}")
        lines.append("")

    body = soup.find("body") or soup

    for elem in body.descendants:
        if elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(elem.name[1])
            text_content = elem.get_text(strip=True)
            if text_content:
                lines.append(f"{'#' * level} {text_content}")
                lines.append("")
        elif elem.name == "p":
            text_content = elem.get_text(strip=True)
            if text_content:
                lines.append(text_content)
                lines.append("")
        elif elem.name == "li":
            text_content = elem.get_text(strip=True)
            if text_content:
                lines.append(f"- {text_content}")
        elif elem.name == "blockquote":
            text_content = elem.get_text(strip=True)
            if text_content:
                for line in text_content.splitlines():
                    lines.append(f"> {line}")
                lines.append("")
        elif elem.name in ("pre", "code"):
            # Only handle if parent is not already a pre/code
            if elem.name == "pre" or (elem.name == "code" and elem.parent.name != "pre"):
                text_content = elem.get_text(strip=False)
                if text_content:
                    lines.append("```")
                    lines.append(text_content)
                    lines.append("```")
                    lines.append("")
        elif elem.name == "table":
            table_lines = extract_html_table(elem)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")

    result_text = "\n".join(lines)
    return result_text, f"bs4+{parser}", warnings


def extract_html_table(table_elem) -> list[str]:
    """Convert HTML table to Markdown table or field-list."""
    rows = []
    for tr in table_elem.find_all("tr"):
        cells = []
        for td in tr.find_all(["td", "th"]):
            text = td.get_text(strip=True).replace("|", "\\|").replace("\n", " ")
            cells.append(text)
        if cells:
            rows.append(cells)

    if not rows:
        return []

    # If table is too wide or complex, convert to field-list
    max_cols = max(len(r) for r in rows)
    if max_cols > 4 or len(rows) > 10:
        lines = []
        for i, row in enumerate(rows):
            if i == 0:
                # Header row
                for j, cell in enumerate(row):
                    lines.append(f"field_{j}: {cell}")
            else:
                for j, cell in enumerate(row):
                    if j < len(row):
                        lines.append(f"  {cell}")
                lines.append("")
        return lines

    # Simple table → Markdown table
    lines = []
    for i, row in enumerate(rows):
        lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            lines.append("|" + "|".join([" --- " for _ in row]) + "|")
    return lines


def extract_docx(docx_path: str) -> tuple[str, str, list[str]]:
    """Extract text from DOCX using python-docx."""
    warnings = []

    # Check if python-docx is available in current interpreter
    try:
        import docx
        return _extract_docx_with_docx(docx_path)
    except ImportError:
        pass

    # Fallback: try isolated venv
    if os.path.exists(VENV_PYTHON):
        try:
            return _extract_docx_via_venv(docx_path)
        except Exception as e:
            warnings.append(f"venv fallback failed: {e}")

    return "", "", warnings + [
        "ERROR: python-docx not installed.",
        "Install in isolated venv:",
        f"  {VENV_PYTHON} -m pip install python-docx",
    ]


def _extract_docx_with_docx(docx_path: str) -> tuple[str, str, list[str]]:
    """Extract using python-docx in current interpreter."""
    from docx import Document

    doc = Document(docx_path)
    lines = []
    warnings = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", ""))
                lines.append(f"{'#' * level} {text}")
            except ValueError:
                lines.append(f"## {text}")
        else:
            lines.append(text)
        lines.append("")

    # Extract tables
    for table in doc.tables:
        table_lines = extract_docx_table(table)
        if table_lines:
            lines.extend(table_lines)
            lines.append("")

    return "\n".join(lines), "python-docx", warnings


def _extract_docx_via_venv(docx_path: str) -> tuple[str, str, list[str]]:
    """Extract using isolated venv python-docx via subprocess."""
    helper_script = """
import json
import sys
from docx import Document

def extract(path):
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", ""))
                lines.append("#" * level + " " + text)
            except ValueError:
                lines.append("## " + text)
        else:
            lines.append(text)
        lines.append("")

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\\\|").replace("\\n", " ") for cell in row.cells]
            if cells:
                rows.append(cells)
        if rows:
            for i, r in enumerate(rows):
                lines.append("| " + " | ".join(r) + " |")
                if i == 0:
                    lines.append("|" + "|".join([" --- " for _ in r]) + "|")
            lines.append("")

    return "\\n".join(lines)

if __name__ == "__main__":
    path = sys.argv[1]
    text = extract(path)
    print(json.dumps({"text": text, "method": "python-docx-venv"}))
"""

    result = subprocess.run(
        [VENV_PYTHON, "-c", helper_script, docx_path],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"venv helper failed: {result.stderr}")

    data = json.loads(result.stdout)
    return data["text"], data["method"], []


def extract_docx_table(table) -> list[str]:
    """Convert python-docx table to Markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
        if cells:
            rows.append(cells)

    if not rows:
        return []

    lines = []
    for i, row in enumerate(rows):
        lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            lines.append("|" + "|".join([" --- " for _ in row]) + "|")
    return lines


def extract_epub(epub_path: str) -> tuple[str, str, list[str]]:
    """Extract text from EPUB using ebooklib."""
    warnings = []

    # Check if ebooklib is available in current interpreter
    try:
        import ebooklib
        from ebooklib import epub
        return _extract_epub_with_ebooklib(epub_path)
    except ImportError:
        pass

    # Fallback: try isolated venv
    if os.path.exists(VENV_PYTHON):
        try:
            return _extract_epub_via_venv(epub_path)
        except Exception as e:
            warnings.append(f"venv fallback failed: {e}")

    return "", "", warnings + [
        "ERROR: ebooklib not installed.",
        "Install in isolated venv:",
        f"  {VENV_PYTHON} -m pip install ebooklib",
    ]


def _extract_epub_with_ebooklib(epub_path: str) -> tuple[str, str, list[str]]:
    """Extract using ebooklib in current interpreter."""
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(epub_path)
    lines = []
    warnings = []

    # Try to get title from metadata
    title = ""
    try:
        titles = book.get_metadata("DC", "title")
        if titles:
            title = titles[0][0]
    except Exception:
        pass

    if title:
        lines.append(f"# {title}")
        lines.append("")

    # Get spine order
    spine_ids = [item[0] for item in book.spine]
    items_by_id = {}
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            items_by_id[item.get_id()] = item

    for item_id in spine_ids:
        if item_id not in items_by_id:
            continue
        item = items_by_id[item_id]
        content = item.get_content().decode("utf-8", errors="replace")

        soup = BeautifulSoup(content, "html.parser")

        # Remove noise elements
        for tag_name in ["script", "style", "noscript", "svg", "nav"]:
            for tag in soup.find_all(tag_name):
                tag.decompose()

        body = soup.find("body") or soup

        # Use find_all with recursive=False for top-level elements only
        # to avoid double-processing nested content
        for elem in body.find_all(recursive=False):
            _process_epub_element(elem, lines, 0)

    return "\n".join(lines), "ebooklib", warnings


def _process_epub_element(elem, lines, depth):
    """Process a single EPUB HTML element recursively."""
    if elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(elem.name[1])
        text_content = elem.get_text(strip=True)
        if text_content:
            lines.append(f"{'#' * level} {text_content}")
            lines.append("")
    elif elem.name == "p":
        text_content = elem.get_text(strip=True)
        if text_content:
            lines.append(text_content)
            lines.append("")
    elif elem.name in ("ul", "ol"):
        _process_epub_list(elem, lines)
    elif elem.name == "blockquote":
        text_content = elem.get_text(strip=True)
        if text_content:
            for line in text_content.splitlines():
                if line.strip():
                    lines.append(f"> {line.strip()}")
            lines.append("")
    elif elem.name == "pre":
        text_content = elem.get_text(strip=False)
        if text_content:
            lines.append("```")
            lines.append(text_content.rstrip())
            lines.append("```")
            lines.append("")
    elif elem.name == "table":
        table_lines = extract_html_table(elem)
        if table_lines:
            lines.extend(table_lines)
            lines.append("")
    elif elem.name in ("div", "section", "article", "header", "footer", "main"):
        # Process container children
        for child in elem.find_all(recursive=False):
            _process_epub_element(child, lines, depth + 1)


def _process_epub_list(list_elem, lines):
    """Process ul/ol element and its li children."""
    for li in list_elem.find_all("li", recursive=False):
        text_content = li.get_text(strip=True)
        if text_content:
            lines.append(f"- {text_content}")
    lines.append("")


def _extract_epub_via_venv(epub_path: str) -> tuple[str, str, list[str]]:
    """Extract using isolated venv ebooklib via subprocess."""
    helper_script = '''
import json
import sys
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

def process_element(elem, lines, depth):
    if elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(elem.name[1])
        text_content = elem.get_text(strip=True)
        if text_content:
            lines.append("#" * level + " " + text_content)
            lines.append("")
    elif elem.name == "p":
        text_content = elem.get_text(strip=True)
        if text_content:
            lines.append(text_content)
            lines.append("")
    elif elem.name in ("ul", "ol"):
        for li in elem.find_all("li", recursive=False):
            text_content = li.get_text(strip=True)
            if text_content:
                lines.append("- " + text_content)
        lines.append("")
    elif elem.name == "blockquote":
        text_content = elem.get_text(strip=True)
        if text_content:
            for line in text_content.splitlines():
                if line.strip():
                    lines.append("> " + line.strip())
            lines.append("")
    elif elem.name == "pre":
        text_content = elem.get_text(strip=False)
        if text_content:
            lines.append("```")
            lines.append(text_content.rstrip())
            lines.append("```")
            lines.append("")
    elif elem.name == "table":
        rows = []
        for tr in elem.find_all("tr"):
            cells = []
            for td in tr.find_all(["td", "th"]):
                text = td.get_text(strip=True).replace("|", "\\\\|").replace("\\n", " ")
                cells.append(text)
            if cells:
                rows.append(cells)
        if rows:
            for i, row in enumerate(rows):
                lines.append("| " + " | ".join(row) + " |")
                if i == 0:
                    lines.append("|" + "|".join([" --- " for _ in row]) + "|")
            lines.append("")
    elif elem.name in ("div", "section", "article", "header", "footer", "main"):
        for child in elem.find_all(recursive=False):
            process_element(child, lines, depth + 1)

def extract(path):
    book = epub.read_epub(path)
    lines = []

    title = ""
    try:
        titles = book.get_metadata("DC", "title")
        if titles:
            title = titles[0][0]
    except Exception:
        pass

    if title:
        lines.append("# " + title)
        lines.append("")

    spine_ids = [item[0] for item in book.spine]
    items_by_id = {}
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            items_by_id[item.get_id()] = item

    for item_id in spine_ids:
        if item_id not in items_by_id:
            continue
        item = items_by_id[item_id]
        content = item.get_content().decode("utf-8", errors="replace")
        soup = BeautifulSoup(content, "html.parser")

        for tag_name in ["script", "style", "noscript", "svg", "nav"]:
            for tag in soup.find_all(tag_name):
                tag.decompose()

        body = soup.find("body") or soup
        for elem in body.find_all(recursive=False):
            process_element(elem, lines, 0)

    return "\\n".join(lines)

if __name__ == "__main__":
    path = sys.argv[1]
    text = extract(path)
    print(json.dumps({"text": text, "method": "ebooklib-venv"}, ensure_ascii=False))
'''

    result = subprocess.run(
        [VENV_PYTHON, "-c", helper_script, epub_path],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"venv helper failed: {result.stderr}")

    data = json.loads(result.stdout)
    return data["text"], data["method"], []


def guess_title(text: str, filename: str) -> str:
    """Guess document title from first heading or filename."""
    lines = text.splitlines()

    # Try meta tags first (for HTML sources)
    # These would have been extracted as # title by extract_html

    # Try first markdown heading
    for line in lines[:50]:
        m = re.match(r"^#\s+(.+)", line.strip())
        if m:
            candidate = m.group(1).strip()
            # Skip filename-like titles
            if not _looks_like_filename(candidate):
                return candidate

    # Try second heading if first was filename-like
    first_heading = None
    for line in lines[:50]:
        m = re.match(r"^#\s+(.+)", line.strip())
        if m:
            candidate = m.group(1).strip()
            if first_heading is None:
                first_heading = candidate
                if not _looks_like_filename(candidate):
                    return candidate
            else:
                # Second heading - use if first was filename-like
                if _looks_like_filename(first_heading) and not _looks_like_filename(candidate):
                    return candidate
                return first_heading

    # Try first non-empty line
    for line in lines[:20]:
        stripped = line.strip()
        if stripped and len(stripped) > 3 and not _looks_like_filename(stripped):
            return stripped

    # Fallback to filename
    return Path(filename).stem.replace("-", " ").replace("_", " ").title()


def _looks_like_filename(text: str) -> bool:
    """Check if text looks like a filename rather than a document title."""
    if not text:
        return True
    # Ends with file extension
    if re.search(r"\.(pdf|html|htm|docx|epub|txt|md)$", text, re.IGNORECASE):
        return True
    # Contains only filename characters (no spaces, mostly alphanumeric + _ -)
    if re.match(r"^[A-Za-z0-9_\-\.]+$", text):
        return True
    # Looks like a stem with underscores (e.g. "art_history")
    if "_" in text and " " not in text and len(text.split("_")) >= 2:
        return True
    return False


def detect_format(input_path: str) -> str:
    """Detect format by extension first, then magic bytes fallback."""
    ext = Path(input_path).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return ext.lstrip(".")

    # Magic bytes fallback for extensionless files
    with open(input_path, "rb") as f:
        header = f.read(8)
    if header[:4] == b"%PDF":
        return "pdf"
    if header[:2] == b"PK":
        # Could be DOCX or EPUB — without extension, we can't tell safely
        return ""

    return ""


def main():
    if len(sys.argv) < 2:
        print(json.dumps({
            "ok": False,
            "errors": ["Usage: extract_text.py <path-to-document> [--mode technical|text]"],
        }))
        sys.exit(1)

    input_path = sys.argv[1]
    mode = "text"
    if "--mode" in sys.argv:
        idx = sys.argv.index("--mode")
        if idx + 1 < len(sys.argv):
            mode = sys.argv[idx + 1].lower()

    if not os.path.exists(input_path):
        print(json.dumps({
            "ok": False,
            "source_path": input_path,
            "errors": [f"File not found: {input_path}"],
        }))
        sys.exit(1)

    document_format = detect_format(input_path)
    if not document_format:
        print(json.dumps({
            "ok": False,
            "source_path": input_path,
            "errors": [f"Unsupported format. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"],
        }))
        sys.exit(1)

    warnings = []
    text = ""
    method = ""

    if document_format in {"txt", "text", "md", "markdown"}:
        text = read_text_file(input_path)
        if text is None:
            print(json.dumps({
                "ok": False,
                "source_path": input_path,
                "format": document_format,
                "errors": ["Could not read text file"],
            }))
            sys.exit(1)
        method = "direct-read"

    elif document_format == "pdf":
        text, method, warnings = extract_pdf(input_path)
        if not text:
            print(json.dumps({
                "ok": False,
                "source_path": input_path,
                "format": document_format,
                "warnings": warnings,
                "errors": ["PDF extraction failed. Install poppler-utils (pdftotext) or pdfminer.six"],
            }))
            sys.exit(1)

    elif document_format in {"html", "htm"}:
        text, method, warnings = extract_html(input_path)
        if not text:
            print(json.dumps({
                "ok": False,
                "source_path": input_path,
                "format": document_format,
                "warnings": warnings,
                "errors": ["HTML extraction failed"],
            }))
            sys.exit(1)

    elif document_format == "docx":
        text, method, warnings = extract_docx(input_path)
        if not text:
            print(json.dumps({
                "ok": False,
                "source_path": input_path,
                "format": document_format,
                "warnings": warnings,
                "errors": ["DOCX extraction failed"],
            }))
            sys.exit(1)

    elif document_format == "epub":
        text, method, warnings = extract_epub(input_path)
        if not text:
            print(json.dumps({
                "ok": False,
                "source_path": input_path,
                "format": document_format,
                "warnings": warnings,
                "errors": ["EPUB extraction failed. Install ebooklib in isolated venv"],
            }))
            sys.exit(1)

    else:
        print(json.dumps({
            "ok": False,
            "source_path": input_path,
            "format": document_format,
            "errors": [f"Format '{document_format}' not yet enabled in this phase"],
        }))
        sys.exit(1)

    title_guess = guess_title(text, input_path)

    result = {
        "ok": True,
        "source_path": str(Path(input_path).resolve()),
        "format": document_format,
        "extraction_method": method,
        "mode": mode,
        "title_guess": title_guess,
        "text": text,
        "char_count": len(text),
        "word_count": len(text.split()),
        "warnings": warnings,
        "errors": [],
    }

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
